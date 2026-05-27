from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX_PATH = DOCS / "informe-pruebas-integracion-ecosense.docx"
PDF_PATH = DOCS / "informe-pruebas-integracion-ecosense.pdf"


CASES = [
    {
        "id": "IT-01",
        "title": "Creacion exitosa de grupo publico",
        "type": "Camino Correcto",
        "integration": "GrupoApplicationService <-> UsuarioRepositoryPort <-> GrupoRepositoryPort <-> EventPublisher",
        "preconditions": (
            "Repositorio de usuarios con el usuario U001 registrado. Repositorio de grupos vacio. "
            "Publicador de eventos RecordingEventPublisher sin mensajes previos."
        ),
        "input": 'crearGrupo(creadorId = "U001", nombre = "EcoLab", descripcion = "Grupo de reciclaje", tipo = PUBLICO)',
        "steps": [
            "El servicio de aplicacion valida que el creador exista en el repositorio de usuarios.",
            "Se consulta el repositorio de grupos para descartar conflicto por nombre duplicado.",
            "Se invoca GrupoService para crear la entidad de dominio con el creador como ADMINISTRADOR.",
            "Se persiste el grupo en el repositorio y se actualiza el grupoId del usuario creador.",
            "Se publica el evento GROUP_CREATED como efecto lateral equivalente a cola o job interno.",
        ],
        "expected": (
            "El resultado es ResultadoCreacion.Exitoso. La BD in-memory contiene el grupo creado, "
            "el usuario U001 queda enlazado al nuevo grupo y la cola registra un unico evento GROUP_CREATED."
        ),
    },
    {
        "id": "IT-02",
        "title": "Rechazo por nombre de grupo duplicado",
        "type": "Camino de Error: 409 Conflict",
        "integration": "GrupoApplicationService <-> GrupoRepositoryPort <-> UsuarioRepositoryPort <-> EventPublisher",
        "preconditions": (
            "Usuario U001 registrado. Repositorio de grupos con un grupo preexistente llamado EcoLab. "
            "Cola de eventos vacia antes de ejecutar el caso."
        ),
        "input": 'crearGrupo("U001", "EcoLab", "Duplicado", PUBLICO)',
        "steps": [
            "El servicio valida la existencia del usuario creador.",
            "Se consulta el repositorio de grupos por nombre antes de crear una nueva entidad.",
            "El repositorio retorna un grupo existente con el mismo nombre.",
            "El flujo se corta antes de persistir cambios o publicar eventos.",
        ],
        "expected": (
            'El resultado es ResultadoCreacion.Error con mensaje "Ya existe un grupo con ese nombre". '
            "La BD mantiene solo el grupo original, el usuario sigue sin grupo asignado y la cola queda vacia."
        ),
    },
    {
        "id": "IT-03",
        "title": "Union exitosa a grupo publico",
        "type": "Camino Correcto",
        "integration": "GrupoApplicationService <-> GrupoService <-> GrupoRepositoryPort <-> UsuarioRepositoryPort <-> EventPublisher",
        "preconditions": (
            "Usuario U002 sin grupo y con 30 puntos. Grupo publico G001 existente con 100 puntos acumulados. "
            "Repositorios in-memory cargados con esos datos."
        ),
        "input": 'unirseAGrupo(usuarioId = "U002", grupoId = "G001")',
        "steps": [
            "El servicio carga usuario y grupo desde los repositorios.",
            "Se reconstruye el estado de dominio en GrupoService con los grupos disponibles.",
            "GrupoService aplica la regla de union a grupo publico.",
            "Se guarda el usuario con grupoId G001 y se actualiza el grupo con el nuevo miembro.",
            "El puntaje del usuario se suma al puntaje total del grupo y se publica GROUP_JOINED.",
        ],
        "expected": (
            "El resultado es ResultadoUnion.Exitoso. El usuario U002 queda asociado a G001, "
            "el grupo contiene al nuevo miembro, el puntaje grupal sube de 100 a 130 y la cola contiene GROUP_JOINED."
        ),
    },
    {
        "id": "IT-04",
        "title": "Intento de union a grupo inexistente",
        "type": "Camino de Error: 404 Not Found",
        "integration": "GrupoApplicationService <-> GrupoRepositoryPort <-> UsuarioRepositoryPort <-> EventPublisher",
        "preconditions": "Usuario U002 registrado. Repositorio de grupos vacio. Cola sin eventos.",
        "input": 'unirseAGrupo("U002", "G404")',
        "steps": [
            "El servicio recupera el usuario desde el repositorio.",
            "Se busca el grupo G404 en el repositorio de grupos.",
            "El repositorio retorna null, por lo que se aborta el flujo de dominio.",
        ],
        "expected": (
            'El resultado es ResultadoUnion.Error con mensaje "El grupo no existe". '
            "El usuario permanece sin grupo, la BD de grupos sigue vacia y no se publican mensajes."
        ),
    },
    {
        "id": "IT-05",
        "title": "Creacion exitosa de solicitud de reciclaje",
        "type": "Camino Correcto",
        "integration": "RecyclingApplicationService <-> ImageStorageClient <-> RecyclingRequestRepositoryPort <-> UsuarioRepositoryPort <-> EventPublisher",
        "preconditions": (
            "Usuario U010 existente. Repositorio de solicitudes vacio. Cliente storage fake configurado para "
            "retornar una URL valida. Cola de eventos vacia."
        ),
        "input": 'submitRequest("U010", "Plastico", 2.5, byteArrayOf(1, 2, 3), "Botellas PET")',
        "steps": [
            "El servicio valida que el usuario exista, que el material no venga vacio y que la cantidad sea positiva.",
            "Se invoca ImageStorageClient para subir la imagen de evidencia.",
            "Con la URL retornada se crea la solicitud REQ-1 en estado PROCESSING.",
            "Se persiste la solicitud en el repositorio y se agrega su id al historial del usuario.",
            "Se publica RECYCLING_VALIDATION_REQUESTED para representar el job posterior de validacion.",
        ],
        "expected": (
            "La solicitud queda persistida con estado PROCESSING, URL de storage, descripcion y cantidad correctas. "
            "El historial del usuario incluye REQ-1 y la cola contiene RECYCLING_VALIDATION_REQUESTED."
        ),
    },
    {
        "id": "IT-06",
        "title": "Timeout del storage al crear solicitud",
        "type": "Camino de Error: Timeout / 500 externo",
        "integration": "RecyclingApplicationService <-> ImageStorageClient <-> RecyclingRequestRepositoryPort <-> UsuarioRepositoryPort <-> EventPublisher",
        "preconditions": (
            "Usuario U010 existente. Repositorio de solicitudes vacio. Cliente storage fake configurado para "
            'lanzar RuntimeException("timeout al subir imagen").'
        ),
        "input": 'submitRequest("U010", "Vidrio", 1.2, byteArrayOf(9), null)',
        "steps": [
            "El servicio valida los datos basicos y llama al cliente externo de storage.",
            "El cliente storage simula un timeout antes de retornar URL.",
            "La excepcion se propaga como Result.failure y se detiene el flujo.",
            "No se persisten solicitudes, no se modifica el historial y no se publica job de validacion.",
        ],
        "expected": (
            'La operacion falla con mensaje "timeout al subir imagen". '
            "El repositorio de solicitudes queda vacio, el historial del usuario no cambia y la cola queda vacia."
        ),
    },
    {
        "id": "IT-07",
        "title": "Canje exitoso de recompensa",
        "type": "Camino Correcto",
        "integration": "RecyclingApplicationService <-> RecyclingRequestRepositoryPort <-> UsuarioRepositoryPort <-> EventPublisher",
        "preconditions": (
            "Usuario U020 con 10 puntos. Solicitud REQ-7 existente, asociada a U020 y en estado REWARD. "
            "Cola de eventos vacia."
        ),
        "input": 'redeemReward(requestId = "REQ-7", userId = "U020", rewardPoints = 25)',
        "steps": [
            "El servicio recupera la solicitud y valida que exista.",
            "Se recupera el usuario y se comprueba que la solicitud le pertenezca.",
            "Se actualiza la solicitud a estado REEDEMED y se registra reward = 25.",
            "Se suman 25 puntos al usuario y se publica REWARD_REDEEMED.",
        ],
        "expected": (
            "La solicitud REQ-7 queda en REEDEMED, el reward queda en 25, "
            "el usuario sube de 10 a 35 puntos y la cola registra REWARD_REDEEMED."
        ),
    },
    {
        "id": "IT-08",
        "title": "Canje de recompensa inexistente",
        "type": "Camino de Error: 404 Not Found",
        "integration": "RecyclingApplicationService <-> RecyclingRequestRepositoryPort <-> UsuarioRepositoryPort <-> EventPublisher",
        "preconditions": "Usuario U020 con 10 puntos. No existe una solicitud REQ-404. Cola vacia.",
        "input": 'redeemReward("REQ-404", "U020", 25)',
        "steps": [
            "El servicio busca la solicitud REQ-404 en el repositorio.",
            "El repositorio retorna null.",
            "El flujo termina como Result.failure antes de modificar puntos o publicar eventos.",
        ],
        "expected": (
            'La operacion falla con mensaje "Solicitud no encontrada". '
            "El usuario mantiene 10 puntos y la cola no recibe eventos."
        ),
    },
]


SUMMARY_ROWS = [
    ["Suite", "Tests", "Fallos", "Errores", "Omitidos", "Resultado"],
    ["Integracion", "15", "0", "0", "0", "PASSED"],
    ["JVM completa", "67", "0", "0", "0", "BUILD SUCCESSFUL"],
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_color(cell, color: RGBColor) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = color


def set_table_borders(table, color: str = "D9E2E8") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("EcoSense - Pruebas de integracion")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(102, 112, 121)


def add_docx_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(34, 92, 87)
        run.font.name = "Aptos Display"


def add_docx_label(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    label_run = paragraph.add_run(label)
    label_run.bold = True
    label_run.font.color.rgb = RGBColor(34, 92, 87)
    paragraph.add_run(value)


def add_steps_table(document: Document, steps: list[str]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    widths = [Cm(2.0), Cm(14.5)]
    for i, text in enumerate(["Paso", "Descripcion"]):
        cell = table.rows[0].cells[i]
        cell.text = text
        cell.width = widths[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "225C57")
        set_cell_text_color(cell, RGBColor(255, 255, 255))
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    for index, step in enumerate(steps, 1):
        row = table.add_row()
        row.cells[0].text = str(index)
        row.cells[1].text = step
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_docx() -> None:
    DOCS.mkdir(exist_ok=True)
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    add_page_number(section)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.12
    styles["Normal"].paragraph_format.space_after = Pt(7)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(140)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Pruebas de software")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(34, 92, 87)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Implementacion pruebas de integracion")
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor(45, 55, 72)

    document.add_paragraph()
    project = document.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.add_run("Nombre del proyecto: ").bold = True
    project.add_run("EcoSense")

    members = document.add_paragraph()
    members.alignment = WD_ALIGN_PARAGRAPH.CENTER
    members.add_run("Integrantes:\n").bold = True
    members.add_run("- Billy Martinez\n- Bastian Lagos")

    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.add_run("Fecha de ejecucion: 27/05/2026")

    document.add_page_break()

    add_docx_heading(document, "Estrategia de implementacion")
    document.add_paragraph(
        "La estrategia de integracion seleccionada para EcoSense es hibrida. "
        "El sistema combina servicios de dominio puros con servicios de aplicacion que coordinan repositorios, "
        "clientes externos y efectos laterales. Por este motivo, se aplica un enfoque bottom-up para validar la "
        "logica de dominio reutilizable y un enfoque top-down para verificar la colaboracion entre los componentes "
        "que ejecutan cada caso de uso."
    )
    document.add_paragraph(
        "Las pruebas se ejecutan a nivel de servicio/modulo. No se prueba la interfaz grafica, no se usa emulador "
        "Android y no se usan webdrivers. Los puntos de entrada son GrupoApplicationService y "
        "RecyclingApplicationService, que representan APIs internas de aplicacion. Firebase, Storage y la cola se "
        "modelan mediante puertos y dobles in-memory controlados por el test."
    )
    document.add_paragraph(
        "Esta decision permite verificar integraciones criticas de forma reproducible: persistencia de usuarios, "
        "grupos y solicitudes; invocacion de un cliente externo de almacenamiento; y publicacion de eventos que "
        "representan jobs o mensajes posteriores de validacion."
    )

    add_docx_heading(document, "Componentes bajo prueba", 2)
    components = document.add_table(rows=1, cols=3)
    components.alignment = WD_TABLE_ALIGNMENT.CENTER
    components.style = "Table Grid"
    set_table_borders(components)
    for i, header in enumerate(["Componente", "Rol", "Doble usado en pruebas"]):
        cell = components.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "225C57")
        set_cell_text_color(cell, RGBColor(255, 255, 255))
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    for row_data in [
        ["GrupoApplicationService", "Orquesta creacion, union y ranking de grupos", "Repositorios in-memory y cola fake"],
        ["RecyclingApplicationService", "Orquesta solicitud de reciclaje y canje de recompensa", "Storage fake, repositorios in-memory y cola fake"],
        ["GrupoService / RankingService", "Reglas puras de dominio", "Instancias reales"],
        ["EventPublisher", "Efecto lateral: cola o job programado", "RecordingEventPublisher"],
    ]:
        row = components.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
            row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_page_break()
    add_docx_heading(document, "Diseno de pruebas de integracion")

    for case in CASES:
        add_docx_heading(document, f"{case['id']} {case['title']} ({case['type']})", 2)
        add_docx_label(document, "Integracion: ", case["integration"])
        add_docx_label(document, "Precondiciones: ", case["preconditions"])
        add_docx_label(document, "Entradas: ", case["input"])
        add_docx_heading(document, "Pasos detallados", 3)
        add_steps_table(document, case["steps"])
        add_docx_heading(document, "Resultados esperados", 3)
        document.add_paragraph(case["expected"])
        if case["id"] in {"IT-02", "IT-04", "IT-06"}:
            document.add_page_break()

    document.add_page_break()
    add_docx_heading(document, "Resultado de los test de integracion")
    document.add_paragraph(
        "Las pruebas se ejecutaron localmente en JVM con Gradle, sin levantar interfaz grafica ni emuladores. "
        "El reporte JUnit XML generado por Gradle confirma cero fallos y cero errores para la suite de integracion."
    )

    summary = document.add_table(rows=1, cols=6)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.style = "Table Grid"
    set_table_borders(summary)
    for i, header in enumerate(SUMMARY_ROWS[0]):
        cell = summary.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "225C57")
        set_cell_text_color(cell, RGBColor(255, 255, 255))
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8)
    for row_data in SUMMARY_ROWS[1:]:
        row = summary.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = value
            row[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in row[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    add_docx_heading(document, "Comandos de ejecucion", 2)
    for command in [
        r"powershell -ExecutionPolicy Bypass -File .\scripts\run-integration-tests.ps1",
        r".\gradlew.bat :app:testDebugUnitTest --tests ""com.ecosense.integration.EcoSenseIntegrationSpec"" --rerun-tasks",
        r".\gradlew.bat :app:testDebugUnitTest",
    ]:
        paragraph = document.add_paragraph()
        paragraph.style = styles["Normal"]
        run = paragraph.add_run(command)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    add_docx_heading(document, "Salida relevante", 2)
    output_lines = [
        '<testsuite name="com.ecosense.integration.EcoSenseIntegrationSpec" tests="15" skipped="0" failures="0" errors="0" timestamp="2026-05-27T05:25:13.592Z" time="0.1">',
        "EcoSenseIntegrationSpec ... IT-01 a IT-08 PASSED",
        "BUILD SUCCESSFUL in 5s",
    ]
    for line in output_lines:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(45, 55, 72)

    add_docx_heading(document, "Conclusion", 2)
    document.add_paragraph(
        "El conjunto de pruebas cubre cuatro flujos criticos del dominio EcoSense y ocho casos de integracion. "
        "Cada flujo valida un camino correcto y un camino de error, comprobando tanto el estado persistido como "
        "los efectos laterales esperados. La ejecucion final fue exitosa, por lo que la evidencia es reproducible "
        "con los comandos indicados."
    )

    document.save(DOCX_PATH)


def style_pdf_table(table: Table, header_rows: int = 1) -> Table:
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, header_rows - 1), colors.HexColor("#225C57")),
                ("TEXTCOLOR", (0, 0), (-1, header_rows - 1), colors.white),
                ("FONTNAME", (0, 0), (-1, header_rows - 1), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (-1, header_rows - 1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9E2E8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def build_pdf() -> None:
    DOCS.mkdir(exist_ok=True)
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=2.0 * cm,
        leftMargin=2.0 * cm,
        topMargin=2.0 * cm,
        bottomMargin=1.8 * cm,
    )
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("CoverTitle", parent=styles["Title"], fontSize=26, textColor=colors.HexColor("#225C57"), alignment=TA_CENTER, spaceAfter=14))
    styles.add(ParagraphStyle("CoverSub", parent=styles["Normal"], fontSize=16, alignment=TA_CENTER, spaceAfter=24))
    styles.add(ParagraphStyle("H1Eco", parent=styles["Heading1"], fontSize=18, textColor=colors.HexColor("#225C57"), spaceAfter=12))
    styles.add(ParagraphStyle("H2Eco", parent=styles["Heading2"], fontSize=13, textColor=colors.HexColor("#225C57"), spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle("BodyEco", parent=styles["BodyText"], fontSize=9.3, leading=12, spaceAfter=7))
    styles.add(ParagraphStyle("CodeEco", parent=styles["Code"], fontSize=7.4, leading=9, textColor=colors.HexColor("#2D3748")))
    story = []

    story.extend([
        Spacer(1, 5.0 * cm),
        Paragraph("Pruebas de software", styles["CoverTitle"]),
        Paragraph("Implementacion pruebas de integracion", styles["CoverSub"]),
        Paragraph("<b>Nombre del proyecto:</b> EcoSense", styles["CoverSub"]),
        Paragraph("<b>Integrantes:</b><br/>- Billy Martinez<br/>- Bastian Lagos", styles["CoverSub"]),
        Paragraph("Fecha de ejecucion: 27/05/2026", styles["CoverSub"]),
        PageBreak(),
    ])

    story.append(Paragraph("Estrategia de implementacion", styles["H1Eco"]))
    for text in [
        "La estrategia de integracion seleccionada para EcoSense es hibrida. El sistema combina servicios de dominio puros con servicios de aplicacion que coordinan repositorios, clientes externos y efectos laterales.",
        "Las pruebas se ejecutan a nivel de servicio/modulo. No se prueba la interfaz grafica, no se usa emulador Android y no se usan webdrivers. Los puntos de entrada son GrupoApplicationService y RecyclingApplicationService.",
        "Firebase, Storage y la cola se modelan mediante puertos y dobles in-memory controlados por el test, permitiendo evidencia reproducible sin red.",
    ]:
        story.append(Paragraph(text, styles["BodyEco"]))

    story.append(Paragraph("Componentes bajo prueba", styles["H2Eco"]))
    component_rows = [
        ["Componente", "Rol", "Doble usado"],
        ["GrupoApplicationService", "Orquesta grupos", "Repositorios in-memory y cola fake"],
        ["RecyclingApplicationService", "Orquesta reciclaje y recompensas", "Storage fake y repositorios in-memory"],
        ["GrupoService / RankingService", "Reglas puras de dominio", "Instancias reales"],
        ["EventPublisher", "Cola o job programado", "RecordingEventPublisher"],
    ]
    story.append(style_pdf_table(Table(component_rows, colWidths=[4.0 * cm, 6.0 * cm, 5.5 * cm])))
    story.append(PageBreak())

    story.append(Paragraph("Diseno de pruebas de integracion", styles["H1Eco"]))
    for case in CASES:
        story.append(Paragraph(f"{case['id']} {case['title']} ({case['type']})", styles["H2Eco"]))
        story.append(Paragraph(f"<b>Integracion:</b> {case['integration']}", styles["BodyEco"]))
        story.append(Paragraph(f"<b>Precondiciones:</b> {case['preconditions']}", styles["BodyEco"]))
        story.append(Paragraph(f"<b>Entradas:</b> {case['input']}", styles["BodyEco"]))
        rows = [["Paso", "Descripcion"]] + [[str(i), step] for i, step in enumerate(case["steps"], 1)]
        story.append(style_pdf_table(Table(rows, colWidths=[1.5 * cm, 14.0 * cm])))
        story.append(Spacer(1, 0.18 * cm))
        story.append(Paragraph(f"<b>Resultados esperados:</b> {case['expected']}", styles["BodyEco"]))
        if case["id"] in {"IT-02", "IT-04", "IT-06"}:
            story.append(PageBreak())

    story.append(PageBreak())
    story.append(Paragraph("Resultado de los test de integracion", styles["H1Eco"]))
    story.append(Paragraph("Las pruebas se ejecutaron localmente en JVM con Gradle, sin interfaz grafica ni emuladores.", styles["BodyEco"]))
    story.append(style_pdf_table(Table(SUMMARY_ROWS, colWidths=[6.0 * cm, 2.0 * cm, 2.0 * cm, 2.0 * cm, 2.0 * cm, 3.0 * cm])))
    story.append(Paragraph("Comandos de ejecucion", styles["H2Eco"]))
    for command in [
        r"powershell -ExecutionPolicy Bypass -File .\scripts\run-integration-tests.ps1",
        r".\gradlew.bat :app:testDebugUnitTest --tests ""com.ecosense.integration.EcoSenseIntegrationSpec"" --rerun-tasks",
        r".\gradlew.bat :app:testDebugUnitTest",
    ]:
        story.append(Paragraph(command, styles["CodeEco"]))
    story.append(Paragraph("Salida relevante", styles["H2Eco"]))
    for line in [
        '&lt;testsuite name="com.ecosense.integration.EcoSenseIntegrationSpec" tests="15" skipped="0" failures="0" errors="0" timestamp="2026-05-27T05:25:13.592Z" time="0.1"&gt;',
        "EcoSenseIntegrationSpec ... IT-01 a IT-08 PASSED",
        "BUILD SUCCESSFUL in 5s",
    ]:
        story.append(Paragraph(line, styles["CodeEco"]))
    story.append(Paragraph("Conclusion", styles["H2Eco"]))
    story.append(Paragraph(
        "El conjunto cubre cuatro flujos criticos del dominio EcoSense y ocho casos de integracion. "
        "Cada flujo valida un camino correcto y un camino de error, comprobando estado persistido y efectos laterales.",
        styles["BodyEco"],
    ))

    doc.build(story)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
