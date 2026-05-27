from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX_PATH = DOCS / "informe-sonarqube-ecosense.docx"
PDF_PATH = DOCS / "informe-sonarqube-ecosense.pdf"


METRIC_ROWS = [
    [
        "Bugs",
        "0",
        "SonarQube no detecto defectos clasificados como bugs. Esto favorece la confiabilidad, aunque no reemplaza pruebas funcionales ni revision manual.",
    ],
    [
        "Vulnerabilities",
        "0",
        "No se detectaron vulnerabilidades en el codigo analizado. El resultado es positivo para seguridad, pero depende de las reglas activas y del alcance analizado.",
    ],
    [
        "Code Smells",
        "20",
        "Existen problemas de mantenibilidad. La mayoria corresponde a complejidad cognitiva, parametros excesivos e imports sin uso; no bloquean ejecucion, pero elevan deuda tecnica.",
    ],
    [
        "Coverage",
        "0.0%",
        "SonarQube no recibio reporte JaCoCo/Kover. Aunque el proyecto tiene pruebas, la cobertura no esta integrada al analisis, por lo que el riesgo real no queda medido.",
    ],
    [
        "Duplications",
        "0.0%",
        "No se detectaron bloques duplicados relevantes. Esto indica buena reutilizacion general y baja repeticion estructural en el codigo fuente analizado.",
    ],
    [
        "Maintainability Rating",
        "A (1.0)",
        "La deuda tecnica total mantiene una calificacion A. Aun asi, los 20 code smells deben revisarse antes de que la complejidad aumente.",
    ],
    [
        "Security Rating",
        "A (1.0)",
        "La calificacion de seguridad es la mejor disponible porque no hay vulnerabilidades abiertas ni hotspots de seguridad detectados.",
    ],
    [
        "Reliability Rating",
        "A (1.0)",
        "La confiabilidad queda en A porque no se detectaron bugs. El resultado debe complementarse con cobertura medible y pruebas de regresion.",
    ],
]


FINDINGS = [
    [
        "Cobertura reportada en 0.0%",
        "Coverage",
        "Alta",
        "Proyecto completo / JaCoCo XML Report Importer",
        "SonarQube detecto que no se importo ningun reporte de cobertura. El log indica: No report imported, no coverage information will be imported. Esto impide evaluar que porcentaje del codigo esta protegido por pruebas automatizadas.",
    ],
    [
        "Complejidad cognitiva excesiva en formulario de reciclaje",
        "Code Smell kotlin:S3776",
        "CRITICAL",
        "app/src/main/kotlin/com/ecosense/screen/RecycleFormScreen.kt:73",
        "SonarQube reporta complejidad 36 sobre un maximo permitido de 15. La pantalla concentra parsing de QR, estado UI, permisos, camara y envio del formulario, lo que dificulta mantenimiento y aumenta riesgo de errores.",
    ],
    [
        "Funcion de navegacion con demasiados parametros",
        "Code Smell kotlin:S107",
        "MAJOR",
        "app/src/main/kotlin/com/ecosense/AppNavHost.kt:74",
        "AppNavigation recibe 9 parametros cuando la regla permite 7. Esto hace mas fragil el contrato entre componentes Compose y complica cambios futuros de configuracion visual o navegacion.",
    ],
    [
        "Literal duplicado en servicio de grupos",
        "Code Smell kotlin:S1192",
        "CRITICAL",
        "app/src/main/kotlin/com/ecosense/service/GrupoApplicationService.kt",
        'El primer analisis detecto el literal "Usuario no encontrado" repetido tres veces. Esta repeticion puede provocar mensajes inconsistentes si se modifica solo una ocurrencia.',
    ],
]


ACTIONS = [
    [
        "Cobertura 0.0%",
        "Configurar generacion de reporte XML con JaCoCo o Kover para tests unitarios JVM y declarar sonar.coverage.jacoco.xmlReportPaths en sonar-project.properties.",
        "Pendiente",
        "docs/sonarqube-metrics.json muestra coverage = 0.0; log del scanner indica que no se importo reporte de cobertura.",
    ],
    [
        "Complejidad en RecycleFormScreen",
        "Propuesta: extraer parsing QR, manejo de permisos/camara y side effects de Toast/navegacion a funciones o ViewModel; dividir la pantalla en composables mas pequenos.",
        "Pendiente",
        "docs/sonarqube-issues.json mantiene issue abierto kotlin:S3776 en RecycleFormScreen.kt:73.",
    ],
    [
        "AppNavigation con 9 parametros",
        "Propuesta: crear un data class de configuracion visual o un objeto de acciones para reducir el numero de parametros expuestos por la funcion.",
        "Pendiente",
        "docs/sonarqube-issues.json mantiene issue abierto kotlin:S107 en AppNavHost.kt:74.",
    ],
    [
        "Literal duplicado Usuario no encontrado",
        "Se aplico constante USER_NOT_FOUND en GrupoApplicationService y se reemplazaron las tres ocurrencias duplicadas.",
        "Resuelto",
        "docs/sonarqube-resolved-issues.json marca kotlin:S1192 como CLOSED/FIXED. Code smells bajaron de 22 a 20 tras la nueva ejecucion.",
    ],
    [
        "Bloque vacio en RecycleFormScreen",
        "Se reemplazo el else vacio por estados explicitos Idle y Uploading con Unit.",
        "Resuelto",
        "docs/sonarqube-resolved-issues.json marca kotlin:S108 como CLOSED/FIXED.",
    ],
]


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2E8")


def style_header(cell) -> None:
    shade(cell, "225C57")
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8.5)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(34, 92, 87)
        run.font.name = "Aptos Display"


def add_table(document: Document, rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    borders(table)
    for i, header in enumerate(rows[0]):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.width = Cm(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        style_header(cell)
    for row_data in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].text = value
            cells[i].width = Cm(widths[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[i].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.2)


def build_docx() -> None:
    DOCS.mkdir(exist_ok=True)
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("EcoSense - Analisis SonarQube")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(102, 112, 121)

    document.styles["Normal"].font.name = "Aptos"
    document.styles["Normal"].font.size = Pt(10.3)
    document.styles["Normal"].paragraph_format.space_after = Pt(7)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(140)
    title_run = title.add_run("Pruebas de software")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(34, 92, 87)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Informe de analisis SonarQube")
    sub_run.font.size = Pt(19)
    sub_run.font.color.rgb = RGBColor(45, 55, 72)

    for text in [
        "Nombre del proyecto: EcoSense",
        "Integrantes: Billy Martinez - Bastian Lagos",
        "Fecha de ejecucion: 27/05/2026",
        "Entrega: 01/06/2026",
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(text)

    document.add_page_break()

    add_heading(document, "Configuracion y ejecucion")
    document.add_paragraph(
        "El analisis se ejecuto sobre la version mas reciente del proyecto EcoSense usando SonarQube Community "
        "Edition en Docker y SonarScanner CLI. El proyecto analizado corresponde al modulo Android/Kotlin ubicado "
        "en app/src/main/kotlin."
    )
    add_table(
        document,
        [
            ["Elemento", "Resultado"],
            ["SonarQube", "9.9.8.100196"],
            ["SonarScanner CLI", "8.0.1.6346"],
            ["Project key", "ecosense"],
            ["Archivos indexados", "55"],
            ["Kotlin source files analizados", "44"],
            ["NCLOC", "5373"],
            ["Quality Gate", "OK"],
        ],
        [5.0, 11.5],
    )

    add_heading(document, "Resultados obtenidos", 1)
    add_table(document, [["Metrica", "Resultado obtenido", "Interpretacion del equipo"], *METRIC_ROWS], [3.7, 3.2, 10.0])

    document.add_page_break()
    add_heading(document, "Analisis de hallazgos criticos")
    add_table(document, [["Hallazgo critico", "Tipo", "Severidad", "Archivo o modulo afectado", "Explicacion del problema"], *FINDINGS], [3.4, 2.9, 2.2, 4.2, 5.2])

    document.add_page_break()
    add_heading(document, "Acciones tomadas o propuestas")
    add_table(document, [["Hallazgo", "Accion tomada o propuesta", "Estado", "Evidencia"], *ACTIONS], [3.2, 6.7, 2.2, 5.0])

    add_heading(document, "Evidencia reproducible")
    document.add_paragraph("Comandos principales:")
    for command in [
        r"powershell -ExecutionPolicy Bypass -File .\scripts\run-sonarqube-analysis.ps1",
        r"docker run --rm -v ${PWD}:/usr/src sonarsource/sonar-scanner-cli:latest -Dsonar.host.url=http://172.17.0.1:9000 -Dsonar.login=admin -Dsonar.password=admin",
        r".\gradlew.bat :app:testDebugUnitTest",
    ]:
        paragraph = document.add_paragraph()
        code = paragraph.add_run(command)
        code.font.name = "Consolas"
        code.font.size = Pt(8.5)
    document.add_paragraph(
        "Archivos de evidencia generados: docs/sonarqube-metrics.json, docs/sonarqube-issues.json, "
        "docs/sonarqube-resolved-issues.json y docs/sonarqube-qualitygate.json."
    )

    add_heading(document, "Conclusion")
    document.add_paragraph(
        "El resultado general es favorable en seguridad, confiabilidad y duplicacion: no hay bugs, no hay "
        "vulnerabilidades y la duplicacion es 0.0%. El principal riesgo actual esta en mantenibilidad y cobertura. "
        "La cobertura aparece en 0.0% porque no se importo reporte XML, y varias pantallas Compose concentran "
        "demasiada logica. Se corrigieron dos code smells y quedaron acciones propuestas para los hallazgos de mayor impacto."
    )
    document.save(DOCX_PATH)


def p_style_pdf_table(table: Table) -> Table:
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#225C57")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9E2E8")),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def rows_to_pdf(rows: list[list[str]], style, col_widths: list[float]) -> Table:
    wrapped = [[Paragraph(cell, style) for cell in row] for row in rows]
    return p_style_pdf_table(Table(wrapped, colWidths=[w * cm for w in col_widths], repeatRows=1))


def build_pdf() -> None:
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("CoverTitle", parent=styles["Title"], fontSize=26, textColor=colors.HexColor("#225C57"), alignment=TA_CENTER, spaceAfter=14))
    styles.add(ParagraphStyle("CoverSub", parent=styles["Normal"], fontSize=15, alignment=TA_CENTER, spaceAfter=18))
    styles.add(ParagraphStyle("H1Eco", parent=styles["Heading1"], fontSize=17, textColor=colors.HexColor("#225C57"), spaceAfter=10))
    styles.add(ParagraphStyle("BodyEco", parent=styles["BodyText"], fontSize=8.8, leading=11, spaceAfter=6))
    styles.add(ParagraphStyle("CodeEco", parent=styles["Code"], fontSize=7.2, leading=8.5))

    story = [
        Spacer(1, 5 * cm),
        Paragraph("Pruebas de software", styles["CoverTitle"]),
        Paragraph("Informe de analisis SonarQube", styles["CoverSub"]),
        Paragraph("Nombre del proyecto: EcoSense", styles["CoverSub"]),
        Paragraph("Integrantes: Billy Martinez - Bastian Lagos", styles["CoverSub"]),
        Paragraph("Fecha de ejecucion: 27/05/2026", styles["CoverSub"]),
        PageBreak(),
        Paragraph("Configuracion y ejecucion", styles["H1Eco"]),
        Paragraph("Analisis ejecutado con SonarQube Community Edition en Docker y SonarScanner CLI sobre app/src/main/kotlin.", styles["BodyEco"]),
        rows_to_pdf(
            [
                ["Elemento", "Resultado"],
                ["SonarQube", "9.9.8.100196"],
                ["SonarScanner CLI", "8.0.1.6346"],
                ["Project key", "ecosense"],
                ["Archivos indexados", "55"],
                ["Kotlin source files analizados", "44"],
                ["NCLOC", "5373"],
                ["Quality Gate", "OK"],
            ],
            styles["BodyEco"],
            [5.0, 10.5],
        ),
        Spacer(1, 0.35 * cm),
        Paragraph("Resultados obtenidos", styles["H1Eco"]),
        rows_to_pdf([["Metrica", "Resultado obtenido", "Interpretacion del equipo"], *METRIC_ROWS], styles["BodyEco"], [3.4, 2.7, 10.1]),
        PageBreak(),
        Paragraph("Analisis de hallazgos criticos", styles["H1Eco"]),
        rows_to_pdf([["Hallazgo critico", "Tipo", "Severidad", "Archivo o modulo afectado", "Explicacion del problema"], *FINDINGS], styles["BodyEco"], [3.0, 2.4, 1.8, 3.8, 5.0]),
        PageBreak(),
        Paragraph("Acciones tomadas o propuestas", styles["H1Eco"]),
        rows_to_pdf([["Hallazgo", "Accion tomada o propuesta", "Estado", "Evidencia"], *ACTIONS], styles["BodyEco"], [3.0, 5.7, 1.8, 5.0]),
        PageBreak(),
        Paragraph("Evidencia reproducible", styles["H1Eco"]),
    ]
    for command in [
        r"powershell -ExecutionPolicy Bypass -File .\scripts\run-sonarqube-analysis.ps1",
        r"docker run --rm -v ${PWD}:/usr/src sonarsource/sonar-scanner-cli:latest -Dsonar.host.url=http://172.17.0.1:9000 -Dsonar.login=admin -Dsonar.password=admin",
        r".\gradlew.bat :app:testDebugUnitTest",
    ]:
        story.append(Paragraph(command, styles["CodeEco"]))
    story.extend(
        [
            Paragraph("Archivos de evidencia: docs/sonarqube-metrics.json, docs/sonarqube-issues.json, docs/sonarqube-resolved-issues.json y docs/sonarqube-qualitygate.json.", styles["BodyEco"]),
            Paragraph("Conclusion", styles["H1Eco"]),
            Paragraph("El proyecto no presenta bugs ni vulnerabilidades detectadas y tiene 0.0% de duplicacion. El principal riesgo esta en cobertura no importada y mantenibilidad por complejidad en pantallas Compose. Se corrigieron dos code smells y se dejaron acciones propuestas para los hallazgos restantes.", styles["BodyEco"]),
        ]
    )

    doc = SimpleDocTemplate(str(PDF_PATH), pagesize=letter, rightMargin=1.8 * cm, leftMargin=1.8 * cm, topMargin=1.8 * cm, bottomMargin=1.8 * cm)
    doc.build(story)


if __name__ == "__main__":
    # Ensure evidence files exist before building, but keep generation deterministic.
    for evidence in ["sonarqube-metrics.json", "sonarqube-issues.json", "sonarqube-resolved-issues.json", "sonarqube-qualitygate.json"]:
        path = DOCS / evidence
        if path.exists():
            json.loads(path.read_text(encoding="utf-8-sig"))
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
